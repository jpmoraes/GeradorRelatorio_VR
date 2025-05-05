import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wav
import speech_recognition as sr
import google.generativeai as genai
from docx import Document
from docx2pdf import convert
from dotenv import load_dotenv
import os

# Carregar as variáveis de ambiente do arquivo .env
load_dotenv()

#Configurar Gemini
api_keyEnv = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_keyEnv)
modelo = genai.GenerativeModel("gemini-2.0-flash")


# Função para gravar áudio
def gravar_audio(nome_arquivo="audio.wav", duracao=5):
    print("Gravando áudio...")
    fs = 44100
    audio = sd.rec(int(duracao * fs), samplerate=fs, channels=1, dtype='int16')
    sd.wait()
    wav.write(nome_arquivo, fs, audio)
    print(f"Áudio gravado: {nome_arquivo}")
    return nome_arquivo


# Função para transcrever áudio
def transcrever_audio(caminho_audio):
    r = sr.Recognizer()
    with sr.AudioFile(caminho_audio) as source:
        audio = r.record(source)
    try:
        texto = r.recognize_google(audio, language="pt-BR")
        return texto
    except sr.UnknownValueError:
        return "Não foi possível entender o áudio."
    except sr.RequestError as e:
        return f"Erro ao conectar ao serviço de reconhecimento: {e}"


# Envia mensagem ao Gemini
def enviar_para_gemini(mensagem):
    print("Gerando resposta com Gemini...")
    mensagem = "Faça uma avaliação simples dos pontos positivos e negativos sobre a resposta dada em uma entrevista de emprego. Retorne apenas o feedback sem formtaçaõ doi texto da seguinte mensagem: " + mensagem
    try:
        resposta = modelo.generate_content(
            mensagem,
            generation_config={
                "max_output_tokens": 100
            }
        )
        return resposta.text
    except Exception as e:
        return f"Erro ao gerar resposta: {e}"


def montar_documento(resposta):
    
    # === Dados para preencher ===
    dados = {
    "Avaliação": "Entrevista",
    "Texto": resposta,
    }    
    # === Caminhos ===
    modelo_path = "modelo.docx"             # modelo com campos tipo ${nome}
    destino_path = "saida_preenchida.docx"
    doc = Document(modelo_path)
    
    # Substituição em parágrafos
    for p in doc.paragraphs:
        for chave, valor in dados.items():
            if f"${{{chave}}}" in p.text:
                p.text = p.text.replace(f"${{{chave}}}", str(valor))

    # Substituição em tabelas (se houver)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in dados.items():
                    if f"${{{chave}}}" in celula.text:
                        celula.text = celula.text.replace(f"${{{chave}}}", str(valor))

    doc.save(destino_path)
    return destino_path

# Função principal
def main():
    arquivo_audio = gravar_audio()
    try:
        texto = transcrever_audio(arquivo_audio)
        print(f"\nTexto transcrito:\n{texto}")

        if "não foi possível entender" in texto.lower() or "erro ao conectar" in texto.lower():
            print("\nErro na transcrição. Encerrando.")
            return

        resposta = enviar_para_gemini(texto)
        preenchido = montar_documento(resposta)
        
        pdf_final = "resultado.pdf"
        convert(preenchido, pdf_final)
        
        print(f"Documento preenchido salvo em: {preenchido}")
        print(f"PDF gerado com sucesso em: {pdf_final}")
    finally:
        if os.path.exists(arquivo_audio):
            os.remove(arquivo_audio)


if __name__ == "__main__":
    main()
