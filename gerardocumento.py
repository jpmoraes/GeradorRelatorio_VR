from docx import Document
from docx2pdf import convert
import os

# === Função para preencher campos no docx existente ===
def preencher_docx(modelo_path, destino_path, dados):
    doc = Document(modelo_path)
   
    # Substituição em parágrafos
    for p in doc.paragraphs:
        for chave, valor in dados.items():
            if f"${{{chave}}}" in p.text:
                p.text = p.text.replace(f"${{{chave}}}", str(valor))
   
    # Substituição em tabelas (se houver)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in dados.items():
                    if f"${{{chave}}}" in celula.text:
                        celula.text = celula.text.replace(f"${{{chave}}}", str(valor))

    doc.save(destino_path)

# === Dados para preencher ===
dados = {
    "Avaliação": "Ana",
    "Texto": "Notebook Dell XPS 13",
}

# === Caminhos ===
modelo = "modelo.docx"             # modelo com campos tipo ${nome}
preenchido = "saida_preenchida.docx"
pdf_final = "saida_preenchida.pdf"

# === Execução ===
preencher_docx(modelo, preenchido, dados)
convert(preenchido, pdf_final)

print(f"Documento preenchido salvo em: {preenchido}")
print(f"PDF gerado com sucesso em: {pdf_final}")